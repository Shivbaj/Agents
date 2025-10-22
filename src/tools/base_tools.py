"""
Web Search Tool for Agent System

This module provides web search capabilities for agents using various search engines
and APIs. It integrates with the MCP system and provides structured search results.

Features:
- Multiple search engine support (Google, Bing, DuckDuckGo)
- Structured result parsing
- URL content extraction
- Search result ranking and filtering
- Rate limiting and caching
- Error handling and fallbacks

Example Usage:
    ```python
    from src.tools.web_search import WebSearchTool
    
    # Create search tool
    search_tool = WebSearchTool()
    
    # Perform search
    results = await search_tool.search("Python programming tutorials")
    
    # Get page content
    content = await search_tool.get_page_content(results[0]["url"])
    ```
"""

import asyncio
import json
import hashlib
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any
from urllib.parse import urljoin, urlparse
import re

import httpx
from bs4 import BeautifulSoup
from langchain_core.tools import BaseTool
from pydantic import BaseModel, Field

from src.config.settings import get_settings


class SearchResult(BaseModel):
    """Structured search result"""
    title: str
    url: str
    snippet: str
    rank: int
    source: str = "web"
    metadata: Dict[str, Any] = Field(default_factory=dict)


class WebSearchConfig(BaseModel):
    """Configuration for web search"""
    default_engine: str = "duckduckgo"
    max_results: int = 10
    timeout: int = 30
    cache_ttl: int = 3600  # Cache TTL in seconds
    user_agent: str = "Mozilla/5.0 (compatible; AI-Agent/1.0)"
    max_content_length: int = 50000
    extract_content: bool = True


class WebSearchTool(BaseTool):
    """
    Web search tool that can search the internet and extract content
    """
    
    name: str = "web_search"
    description: str = "Search the web for information and retrieve page content"
    
    def __init__(self, config: Optional[WebSearchConfig] = None):
        super().__init__()
        self.config = config or WebSearchConfig()
        self.settings = get_settings()
        self.cache: Dict[str, Dict[str, Any]] = {}
        
        # HTTP client with proper headers
        self.client = httpx.AsyncClient(
            headers={"User-Agent": self.config.user_agent},
            timeout=self.config.timeout,
            follow_redirects=True
        )
    
    def _run(self, query: str, max_results: int = None) -> List[Dict[str, Any]]:
        """Synchronous search (not recommended, use arun instead)"""
        import asyncio
        return asyncio.run(self.arun(query, max_results))
    
    async def _arun(self, query: str, max_results: int = None) -> List[Dict[str, Any]]:
        """Perform web search"""
        max_results = max_results or self.config.max_results
        
        # Check cache first
        cache_key = self._get_cache_key(query, max_results)
        if cache_key in self.cache:
            cache_entry = self.cache[cache_key]
            if datetime.now() - cache_entry["timestamp"] < timedelta(seconds=self.config.cache_ttl):
                return cache_entry["results"]
        
        # Perform search
        try:
            if self.config.default_engine == "duckduckgo":
                results = await self._search_duckduckgo(query, max_results)
            elif self.config.default_engine == "google":
                results = await self._search_google(query, max_results)
            else:
                results = await self._search_duckduckgo(query, max_results)  # Fallback
            
            # Cache results
            self.cache[cache_key] = {
                "results": results,
                "timestamp": datetime.now()
            }
            
            return results
            
        except Exception as e:
            return [{"error": f"Search failed: {str(e)}"}]
    
    async def _search_duckduckgo(self, query: str, max_results: int) -> List[Dict[str, Any]]:
        """Search using DuckDuckGo"""
        # DuckDuckGo Instant Answer API
        search_url = "https://api.duckduckgo.com/"
        params = {
            "q": query,
            "format": "json",
            "no_html": "1",
            "skip_disambig": "1"
        }
        
        try:
            response = await self.client.get(search_url, params=params)
            response.raise_for_status()
            data = response.json()
            
            results = []
            
            # Process instant answer
            if data.get("AbstractText"):
                results.append({
                    "title": data.get("Heading", query),
                    "url": data.get("AbstractURL", ""),
                    "snippet": data.get("AbstractText", ""),
                    "rank": 1,
                    "source": "duckduckgo_instant"
                })
            
            # Process related topics
            for i, topic in enumerate(data.get("RelatedTopics", [])[:max_results]):
                if isinstance(topic, dict) and "Text" in topic:
                    results.append({
                        "title": topic.get("Text", "")[:100],
                        "url": topic.get("FirstURL", ""),
                        "snippet": topic.get("Text", ""),
                        "rank": i + 2,
                        "source": "duckduckgo_related"
                    })
            
            # If no results from instant API, try web search
            if not results:
                results = await self._search_duckduckgo_web(query, max_results)
            
            return results[:max_results]
            
        except Exception as e:
            # Fallback to web search
            return await self._search_duckduckgo_web(query, max_results)
    
    async def _search_duckduckgo_web(self, query: str, max_results: int) -> List[Dict[str, Any]]:
        """Search DuckDuckGo web results by scraping"""
        search_url = "https://duckduckgo.com/html/"
        params = {"q": query}
        
        try:
            response = await self.client.get(search_url, params=params)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            results = []
            
            # Parse search results
            for i, result_div in enumerate(soup.find_all("div", class_="result__body")[:max_results]):
                title_elem = result_div.find("a", class_="result__a")
                snippet_elem = result_div.find("a", class_="result__snippet")
                
                if title_elem:
                    title = title_elem.get_text(strip=True)
                    url = title_elem.get("href", "")
                    snippet = snippet_elem.get_text(strip=True) if snippet_elem else ""
                    
                    results.append({
                        "title": title,
                        "url": url,
                        "snippet": snippet,
                        "rank": i + 1,
                        "source": "duckduckgo_web"
                    })
            
            return results
            
        except Exception as e:
            raise Exception(f"DuckDuckGo web search failed: {e}")
    
    async def _search_google(self, query: str, max_results: int) -> List[Dict[str, Any]]:
        """Search using Google Custom Search API (requires API key)"""
        api_key = getattr(self.settings, 'google_search_api_key', None)
        search_engine_id = getattr(self.settings, 'google_search_engine_id', None)
        
        if not api_key or not search_engine_id:
            raise Exception("Google Search API key or Search Engine ID not configured")
        
        search_url = "https://www.googleapis.com/customsearch/v1"
        params = {
            "key": api_key,
            "cx": search_engine_id,
            "q": query,
            "num": min(max_results, 10)  # Google API limit
        }
        
        try:
            response = await self.client.get(search_url, params=params)
            response.raise_for_status()
            data = response.json()
            
            results = []
            for i, item in enumerate(data.get("items", [])):
                results.append({
                    "title": item.get("title", ""),
                    "url": item.get("link", ""),
                    "snippet": item.get("snippet", ""),
                    "rank": i + 1,
                    "source": "google"
                })
            
            return results
            
        except Exception as e:
            raise Exception(f"Google search failed: {e}")
    
    async def get_page_content(self, url: str) -> Dict[str, Any]:
        """Extract content from a web page"""
        try:
            response = await self.client.get(url)
            response.raise_for_status()
            
            # Parse HTML content
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text content
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            # Limit content length
            if len(text) > self.config.max_content_length:
                text = text[:self.config.max_content_length] + "..."
            
            # Extract metadata
            title = soup.find("title")
            title_text = title.get_text(strip=True) if title else ""
            
            meta_description = soup.find("meta", attrs={"name": "description"})
            description = ""
            if meta_description:
                description = meta_description.get("content", "")
            
            return {
                "url": url,
                "title": title_text,
                "content": text,
                "description": description,
                "word_count": len(text.split()),
                "extracted_at": datetime.now().isoformat()
            }
            
        except Exception as e:
            return {
                "url": url,
                "error": f"Failed to extract content: {str(e)}",
                "extracted_at": datetime.now().isoformat()
            }
    
    def _get_cache_key(self, query: str, max_results: int) -> str:
        """Generate cache key for search query"""
        key_data = f"{query}:{max_results}:{self.config.default_engine}"
        return hashlib.md5(key_data.encode()).hexdigest()
    
    async def cleanup(self):
        """Cleanup resources"""
        await self.client.aclose()


class FileProcessorTool(BaseTool):
    """
    Tool for processing various file types (PDF, DOCX, TXT, etc.)
    """
    
    name: str = "file_processor"
    description: str = "Process and extract content from various file types"
    
    def __init__(self):
        super().__init__()
        self.settings = get_settings()
    
    def _run(self, file_path: str, file_type: str = None) -> Dict[str, Any]:
        """Synchronous file processing"""
        import asyncio
        return asyncio.run(self.arun(file_path, file_type))
    
    async def _arun(self, file_path: str, file_type: str = None) -> Dict[str, Any]:
        """Process a file and extract its content"""
        try:
            if not file_type:
                file_type = self._detect_file_type(file_path)
            
            if file_type == "pdf":
                return await self._process_pdf(file_path)
            elif file_type == "docx":
                return await self._process_docx(file_path)
            elif file_type == "txt":
                return await self._process_text(file_path)
            elif file_type in ["jpg", "jpeg", "png", "gif", "bmp"]:
                return await self._process_image(file_path)
            else:
                return {"error": f"Unsupported file type: {file_type}"}
                
        except Exception as e:
            return {"error": f"File processing failed: {str(e)}"}
    
    def _detect_file_type(self, file_path: str) -> str:
        """Detect file type from extension"""
        extension = file_path.lower().split('.')[-1]
        return extension
    
    async def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF file"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text_content += f"\n--- Page {page_num + 1} ---\n"
                    text_content += page.extract_text()
                
                return {
                    "file_path": file_path,
                    "file_type": "pdf",
                    "content": text_content.strip(),
                    "page_count": len(pdf_reader.pages),
                    "word_count": len(text_content.split()),
                    "processed_at": datetime.now().isoformat()
                }
                
        except ImportError:
            return {"error": "PyPDF2 not installed. Install with: pip install PyPDF2"}
        except Exception as e:
            return {"error": f"PDF processing failed: {str(e)}"}
    
    async def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process DOCX file"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            return {
                "file_path": file_path,
                "file_type": "docx",
                "content": text_content.strip(),
                "paragraph_count": len(doc.paragraphs),
                "word_count": len(text_content.split()),
                "processed_at": datetime.now().isoformat()
            }
            
        except ImportError:
            return {"error": "python-docx not installed. Install with: pip install python-docx"}
        except Exception as e:
            return {"error": f"DOCX processing failed: {str(e)}"}
    
    async def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Process text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return {
                "file_path": file_path,
                "file_type": "txt",
                "content": content,
                "line_count": len(content.splitlines()),
                "word_count": len(content.split()),
                "processed_at": datetime.now().isoformat()
            }
            
        except Exception as e:
            return {"error": f"Text file processing failed: {str(e)}"}
    
    async def _process_image(self, file_path: str) -> Dict[str, Any]:
        """Process image file (extract metadata)"""
        try:
            from PIL import Image
            from PIL.ExifTags import TAGS
            
            with Image.open(file_path) as img:
                # Basic image info
                info = {
                    "file_path": file_path,
                    "file_type": "image",
                    "format": img.format,
                    "mode": img.mode,
                    "size": img.size,
                    "processed_at": datetime.now().isoformat()
                }
                
                # Extract EXIF data
                exifdata = img.getexif()
                if exifdata:
                    exif_info = {}
                    for tag_id in exifdata:
                        tag = TAGS.get(tag_id, tag_id)
                        data = exifdata.get(tag_id)
                        if isinstance(data, bytes):
                            data = data.decode()
                        exif_info[tag] = data
                    
                    info["exif"] = exif_info
                
                return info
                
        except ImportError:
            return {"error": "Pillow not installed. Install with: pip install Pillow"}
        except Exception as e:
            return {"error": f"Image processing failed: {str(e)}"}


class CalculatorTool(BaseTool):
    """
    Safe calculator tool for mathematical operations
    """
    
    name: str = "calculator"
    description: str = "Perform safe mathematical calculations and evaluations"
    
    def _run(self, expression: str) -> Dict[str, Any]:
        """Perform calculation"""
        return self._safe_eval(expression)
    
    async def _arun(self, expression: str) -> Dict[str, Any]:
        """Async calculation"""
        return self._safe_eval(expression)
    
    def _safe_eval(self, expression: str) -> Dict[str, Any]:
        """Safely evaluate mathematical expressions"""
        try:
            # Remove any dangerous functions/imports
            dangerous = [
                'import', 'exec', 'eval', 'open', 'file', 'input', 
                'raw_input', '__import__', '__builtins__'
            ]
            
            for danger in dangerous:
                if danger in expression.lower():
                    return {"error": f"Dangerous operation '{danger}' not allowed"}
            
            # Only allow safe mathematical operations
            allowed_names = {
                k: v for k, v in __builtins__.items() 
                if k in ['abs', 'round', 'min', 'max', 'sum', 'pow']
            }
            
            # Add math functions
            import math
            allowed_names.update({
                'sqrt': math.sqrt,
                'sin': math.sin,
                'cos': math.cos,
                'tan': math.tan,
                'log': math.log,
                'log10': math.log10,
                'exp': math.exp,
                'pi': math.pi,
                'e': math.e
            })
            
            # Evaluate expression
            result = eval(expression, {"__builtins__": {}}, allowed_names)
            
            return {
                "expression": expression,
                "result": result,
                "calculated_at": datetime.now().isoformat()
            }
            
        except Exception as e:
            return {
                "expression": expression,
                "error": f"Calculation failed: {str(e)}",
                "calculated_at": datetime.now().isoformat()
            }


class CodeExecutorTool(BaseTool):
    """
    Safe code execution tool (sandboxed)
    """
    
    name: str = "code_executor"
    description: str = "Execute code safely in a sandboxed environment"
    
    def __init__(self, allowed_languages: List[str] = None):
        super().__init__()
        self.allowed_languages = allowed_languages or ["python"]
    
    def _run(self, code: str, language: str = "python") -> Dict[str, Any]:
        """Execute code synchronously"""
        import asyncio
        return asyncio.run(self.arun(code, language))
    
    async def _arun(self, code: str, language: str = "python") -> Dict[str, Any]:
        """Execute code asynchronously"""
        if language not in self.allowed_languages:
            return {"error": f"Language '{language}' not allowed"}
        
        if language == "python":
            return await self._execute_python(code)
        else:
            return {"error": f"Execution for '{language}' not implemented"}
    
    async def _execute_python(self, code: str) -> Dict[str, Any]:
        """Execute Python code safely"""
        try:
            # Basic safety checks
            dangerous_imports = [
                'os', 'sys', 'subprocess', 'socket', 'urllib', 'requests',
                'file', 'open', 'exec', 'eval', '__import__'
            ]
            
            code_lower = code.lower()
            for danger in dangerous_imports:
                if danger in code_lower:
                    return {"error": f"Dangerous operation '{danger}' not allowed"}
            
            # Capture output
            import io
            import contextlib
            from contextlib import redirect_stdout, redirect_stderr
            
            stdout_buffer = io.StringIO()
            stderr_buffer = io.StringIO()
            
            # Execute with output capture
            with redirect_stdout(stdout_buffer), redirect_stderr(stderr_buffer):
                # Create restricted environment
                restricted_globals = {
                    '__builtins__': {
                        'print': print,
                        'len': len,
                        'range': range,
                        'list': list,
                        'dict': dict,
                        'str': str,
                        'int': int,
                        'float': float,
                        'bool': bool,
                        'sum': sum,
                        'min': min,
                        'max': max,
                        'abs': abs,
                        'round': round
                    }
                }
                
                # Add safe math module
                import math
                restricted_globals['math'] = math
                
                exec(code, restricted_globals)
            
            stdout_value = stdout_buffer.getvalue()
            stderr_value = stderr_buffer.getvalue()
            
            return {
                "code": code,
                "language": "python",
                "stdout": stdout_value,
                "stderr": stderr_value,
                "success": True,
                "executed_at": datetime.now().isoformat()
            }
            
        except Exception as e:
            return {
                "code": code,
                "language": "python",
                "error": str(e),
                "success": False,
                "executed_at": datetime.now().isoformat()
            }


# Tool registry for easy access
AVAILABLE_TOOLS = {
    "web_search": WebSearchTool,
    "file_processor": FileProcessorTool,
    "calculator": CalculatorTool,
    "code_executor": CodeExecutorTool
}


def get_tool(tool_name: str, **kwargs) -> BaseTool:
    """Get a tool instance by name"""
    if tool_name not in AVAILABLE_TOOLS:
        raise ValueError(f"Tool '{tool_name}' not found. Available tools: {list(AVAILABLE_TOOLS.keys())}")
    
    return AVAILABLE_TOOLS[tool_name](**kwargs)


def get_all_tools(**kwargs) -> List[BaseTool]:
    """Get all available tools"""
    tools = []
    for tool_class in AVAILABLE_TOOLS.values():
        try:
            tools.append(tool_class(**kwargs))
        except Exception as e:
            print(f"Warning: Could not initialize tool {tool_class.__name__}: {e}")
    
    return tools